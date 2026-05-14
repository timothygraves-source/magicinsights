"""Generates a CI findings .docx matching the Agentforce Operations World Tour NYC style."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import json
import re

NAVY = RGBColor(0x02, 0x2D, 0x60)
BLUE = RGBColor(0x0C, 0x9D, 0xDA)
BODY = RGBColor(0x13, 0x13, 0x14)

FONT_NAME = "Salesforce Sans"


def _set_font(run, size_pt, color, bold=False, italic=False):
    run.font.name = FONT_NAME
    run.font.size = Pt(size_pt)
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic


def _set_spacing(para, before_pt=0, after_pt=0):
    pf = para.paragraph_format
    if before_pt:
        pf.space_before = Pt(before_pt)
    if after_pt:
        pf.space_after = Pt(after_pt)


def _add_title(doc, title_text, subtitle_text, event_text):
    p = doc.add_paragraph()
    run = p.add_run(title_text)
    _set_font(run, 26, NAVY)
    _set_spacing(p, after_pt=6)

    p2 = doc.add_paragraph()
    run2 = p2.add_run(subtitle_text)
    _set_font(run2, 12, BLUE)
    _set_spacing(p2, after_pt=6)

    p3 = doc.add_paragraph()
    run3a = p3.add_run("Salesforce Marketing — Customer Insights Team")
    _set_font(run3a, 9, BODY, italic=True)
    br = OxmlElement("w:br")
    run3a._r.append(br)
    run3b = p3.add_run("This document was created with the help of generative AI tools and edited by humans.")
    _set_font(run3b, 9, BODY, italic=True)
    _set_spacing(p3, after_pt=18)


def _add_section_header(doc, text, space_before_pt=18):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_font(run, 14, BLUE)
    _set_spacing(p, before_pt=space_before_pt, after_pt=8)
    return p


def _add_sub_header(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_font(run, 11, NAVY, bold=True)
    _set_spacing(p, before_pt=9, after_pt=4)
    return p


def _add_body(doc, text, space_after_pt=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_font(run, 10, BODY)
    _set_spacing(p, after_pt=space_after_pt)
    return p


def _add_bullet(doc, label, body_text):
    """Blue bullet + bold navy label + regular body text."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Inches(0.25)
    _set_spacing(p, after_pt=4)

    bullet_run = p.add_run("• ")
    _set_font(bullet_run, 10, BLUE, bold=True)

    if label:
        label_run = p.add_run(label + ": ")
        _set_font(label_run, 10, NAVY, bold=True)

    body_run = p.add_run(body_text)
    _set_font(body_run, 10, BODY)
    return p


def _add_quote(doc, quote_text, attribution):
    """Indented italic block quote with navy attribution on new line."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Inches(0.5)
    _set_spacing(p, before_pt=6, after_pt=6)

    quote_run = p.add_run(f'"{quote_text}"')
    _set_font(quote_run, 10, BODY, italic=True)

    br = OxmlElement("w:br")
    quote_run._r.append(br)

    attr_run = p.add_run(f"— {attribution}")
    _set_font(attr_run, 9, NAVY)
    return p


def _add_toc(doc, sections):
    _add_section_header(doc, "Table of Contents", space_before_pt=12)
    for entry in sections:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.left_indent = Inches(0.2)
        _set_spacing(p, after_pt=4)
        run = p.add_run(entry)
        _set_font(run, 10, BODY)


def generate_docx(findings_json: dict, output_path: str):
    """
    findings_json structure:
    {
      "title": "...",
      "subtitle": "...",          # e.g. "Product Name: Customer Insights"
      "event": "...",             # e.g. "Dreamforce | September 2025"
      "executive_summary": ["paragraph1", "paragraph2", ...],
      "methodology": {
        "description": "...",
        "participants": [
          {"name": "...", "role": "..."},
          ...
        ]
      },
      "key_findings": [
        {
          "title": "Key Finding 1: ...",
          "body": "...",
          "quotes": [
            {"text": "...", "attribution": "..."},
            ...
          ],
          "closing": "..."   # optional paragraph after quotes
        },
        ...
      ],
      "additional_sections": [    # optional extra sections like "Reception of Demo"
        {
          "title": "...",
          "intro": "...",
          "sub_sections": [       # optional
            {"title": "...", "body": "...", "quote": {"text": "...", "attribution": "..."}}
          ],
          "bullets": [            # optional
            {"label": "...", "text": "..."}
          ]
        }
      ],
      "outstanding_questions": [
        {"label": "...", "text": "..."},
        ...
      ],
      "future_requests": {
        "intro": "...",
        "items": ["...", "..."],
        "closing": "..."
      },
      "sources": "..."
    }
    """
    doc = Document()

    # Page margins: 1 inch all sides
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    data = findings_json

    # Title block
    _add_title(doc, data.get("title", "Customer Insights"),
               data.get("subtitle", ""),
               data.get("event", ""))

    # Build TOC entries
    toc_entries = ["Executive Summary", "Methodology & Participants"]
    for kf in data.get("key_findings", []):
        toc_entries.append(kf["title"])
    for sec in data.get("additional_sections", []):
        toc_entries.append(sec["title"])
    toc_entries += ["Outstanding Questions & Gaps",
                    "Requests for Future Use Cases & Product Development",
                    "Sources"]
    _add_toc(doc, toc_entries)

    # Executive Summary
    _add_section_header(doc, "Executive Summary")
    for para in data.get("executive_summary", []):
        _add_body(doc, para)

    # Methodology & Participants
    _add_section_header(doc, "Methodology & Participants")
    _add_body(doc, data["methodology"]["description"])
    _add_body(doc, "Participant profiles:")
    for p in data["methodology"]["participants"]:
        _add_bullet(doc, p["name"], p["role"])

    # Key Findings
    for kf in data.get("key_findings", []):
        _add_section_header(doc, kf["title"])
        _add_body(doc, kf["body"])
        for q in kf.get("quotes", []):
            _add_quote(doc, q["text"], q["attribution"])
        if kf.get("closing"):
            _add_body(doc, kf["closing"])

    # Additional sections (e.g. Reception of Demo)
    for sec in data.get("additional_sections", []):
        _add_section_header(doc, sec["title"])
        if sec.get("intro"):
            _add_body(doc, sec["intro"])
        for sub in sec.get("sub_sections", []):
            _add_sub_header(doc, sub["title"])
            if sub.get("body"):
                _add_body(doc, sub["body"])
            if sub.get("quote"):
                _add_quote(doc, sub["quote"]["text"], sub["quote"]["attribution"])
        for b in sec.get("bullets", []):
            _add_bullet(doc, b.get("label", ""), b.get("text", ""))

    # Outstanding Questions & Gaps
    _add_section_header(doc, "Outstanding Questions & Gaps")
    _add_body(doc, "The following questions and concerns were consistently raised across interviews:")
    for item in data.get("outstanding_questions", []):
        _add_bullet(doc, item.get("label", ""), item.get("text", ""))

    # Future Requests
    _add_section_header(doc, "Requests for Future Use Cases & Product Development")
    fr = data.get("future_requests", {})
    if fr.get("intro"):
        _add_body(doc, fr["intro"])
    for item in fr.get("items", []):
        _add_bullet(doc, "", item)
    if fr.get("closing"):
        _add_body(doc, fr["closing"])

    # Sources
    _add_section_header(doc, "Sources")
    _add_body(doc, data.get("sources", ""))

    doc.save(output_path)
    return output_path
